import pickle
import numpy as np
##import array_to_latex as a2l
import time
import os
import sys
from canvas import EditCanvas
from docx import Document
from docx.shared import Cm, Inches
##import matplotlib
##matplotlib.rcParams['text.usetex'] = True
##from matplotlib import mathtext
import pygame as pg
DIR = 'archives\\'
def cplxrepr(c):
    r, i = c.real,c.imag
    if i < 0:
        return '%.2f%.2fi'%(r, i)
    return '%.2f+%.2fi'%(r, i)
##    return repr(c).removeprefix('(').removesuffix(')').replace('j', 'i')#'%.2f + %.2fi'%(c.real, c.imag)
class Archive:
    def __init__(self, path = None):
        self.path = path
        if not os.path.exists(DIR):
            os.makedirs(DIR)
    def newPath(self, suf = '.pkl'):
        path = DIR + time.strftime(
                '%Y%m%d%H%M%S',
                time.localtime(time.time())
            ) + suf
        return path
    def save(self, canvas):
        if self.path is None:
            self.path = self.newPath()
        with open(self.path, 'wb') as f:
            print((canvas.gvectors, canvas.gpoints, canvas.ref))
            pickle.dump((canvas.gvectors, canvas.gpoints, canvas.ref), f)
    def load(self):
        if self.path is None:
            return (set(), set(), None)
        with open(path, 'rb') as f:
            l = pickle.load(f)
            if len(l) == 2:
                gvectors, gpoints = l
                ref = None
            else:
                gvectors, gpoints, ref = l
        print(gvectors, gpoints, ref)
        return gvectors, gpoints, ref
##    def report(self, law):
##        ss = []
##        for i, op in enumerate(law.ops):
##            s = r"$M(\varphi_%d) = \begin{bmatrix}"%(i+1) + cplxrepr(op[0][0]) + ' & ' + cplxrepr(op[0][1]) + r'\\' + \
##                cplxrepr(op[1][0]) + r' & ' + cplxrepr(op[1][1]) + r'\end{bmatrix}$'
##            ss.append(s)
##        res = '\n\n'.join(ss)
##        res += '\n$d_H = %.3f$\n'%law.dim()
##        path = self.newPath('.md')
##        with open(path, 'w') as f:
##            f.write(res)
    def report(self):
##        print(editCanvas)
        doc = Document()
        pg.image.save(editCanvas.surface, 'law_pic.png')
        pg.image.save(effectCanvas.surface, 'fractal_pic.png')
        doc.add_picture('law_pic.png', width = Inches(5))
        doc.add_picture('fractal_pic.png', width = Inches(5))
        #
        for i, op in enumerate(editCanvas.law.ops):
            s = "M(φ%d) = [\n\t"%(i+1) + cplxrepr(op[0][0]) + '\t' + \
                cplxrepr(op[0][1]) + '\n\t' + cplxrepr(op[1][0]) + '\t' + cplxrepr(op[1][1]) + '\n]'
##            s = r"$M(\varphi_%d) = \begin{bmatrix}"%(i+1) + cplxrepr(op[0][0]) + ' & ' + cplxrepr(op[0][1]) + r'\\' + \
##                cplxrepr(op[1][0]) + r' & ' + cplxrepr(op[1][1]) + r'\end{bmatrix}$'
            doc.add_paragraph(s)
##            img = mathtext.math_to_image(s, 'formula.png')
##            doc.add_picture('formula.png', width = Inches(5))
##        #
        doc.save(self.newPath('.docx'))
        
##print(sys.argv)
if len(sys.argv) == 2:
    path = sys.argv[1]
else:
    path = None
archive = Archive(path)
